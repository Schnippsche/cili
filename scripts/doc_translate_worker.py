#!/usr/bin/env python3
# scripts/doc_translate_worker.py
"""
Übersetzt den Textinhalt eines Dokuments mittels NLLB.

Unterstützte Eingabeformate:
  - .txt  — wird direkt gelesen
  - .pdf  — Textblöcke werden via PyMuPDF extrahiert
  - .docx/.odt/.doc/.pptx/.ppt/.rtf/… — via LibreOffice headless nach TXT konvertiert

Aufruf:
    python doc_translate_worker.py \
        --input  /path/to/doc.pdf \
        --output /path/to/out.txt \
        --source de \
        --target en \
        --model  /opt/cili/nllb/nllb-600M \
        --libreoffice /usr/bin/soffice \
        --device cpu \
        --compute-type int8 \
        --beam-size 4 \
        --max-batch-size 16

Exit-Codes: 0 = OK, 1 = Fehler
"""
import argparse
import os
import re
import subprocess
import sys
import tempfile

# Windows: DLL-Verzeichnisse der nvidia-pip-Pakete registrieren, damit ctranslate2
# cublas64_12.dll usw. findet. Muss VOR jedem ctranslate2-Import passieren.
if sys.platform == "win32":
    import glob as _glob
    import site as _site_mod
    for _site in _site_mod.getsitepackages():
        _nvidia = os.path.join(_site, "nvidia")
        if not os.path.isdir(_nvidia):
            continue
        for _sub in os.listdir(_nvidia):
            _bin = os.path.join(_nvidia, _sub, "bin")
            if os.path.isdir(_bin):
                os.environ["PATH"] = _bin + os.pathsep + os.environ.get("PATH", "")
                try:
                    os.add_dll_directory(_bin)
                except OSError:
                    pass
    for _cuda_bin in _glob.glob(
        r"C:\Program Files\NVIDIA GPU Computing Toolkit\CUDA\v*\bin"
    ):
        os.environ["PATH"] = _cuda_bin + os.pathsep + os.environ.get("PATH", "")
        try:
            os.add_dll_directory(_cuda_bin)
        except OSError:
            pass

import fitz  # PyMuPDF

try:
    from docx import Document as _DocxDocument
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False


def gpu_temperature() -> str:
    """Liest die aktuelle GPU-Temperatur über nvidia-smi aus (z.B. '75'). Liefert '?' falls nicht verfügbar."""
    try:
        out = subprocess.run(
            ['nvidia-smi', '--query-gpu=temperature.gpu', '--format=csv,noheader,nounits'],
            capture_output=True, text=True, timeout=5, check=True,
        )
        return out.stdout.strip().splitlines()[0]
    except Exception:
        return '?'


# ISO 639-1 → NLLB Flores-200 (Sprachcode-Zuordnung)
LANG_MAP = {
    "af": "afr_Latn", "sq": "als_Latn", "hy": "hye_Armn", "ast": "ast_Latn",
    "az": "azj_Latn", "eu": "eus_Latn", "be": "bel_Cyrl", "bs": "bos_Latn",
    "br": "bre_Latn", "bg": "bul_Cyrl", "ca": "cat_Latn", "hr": "hrv_Latn",
    "cs": "ces_Latn", "da": "dan_Latn", "nl": "nld_Latn", "en": "eng_Latn",
    "et": "est_Latn", "fo": "fao_Latn", "fi": "fin_Latn", "fr": "fra_Latn",
    "gl": "glg_Latn", "ka": "kat_Geor", "de": "deu_Latn", "el": "ell_Grek",
    "hu": "hun_Latn", "is": "isl_Latn", "ga": "gle_Latn", "it": "ita_Latn",
    "ja": "jpn_Jpan", "lv": "lav_Latn", "lb": "ltz_Latn", "lt": "lit_Latn",
    "mk": "mkd_Cyrl", "mt": "mlt_Latn", "nb": "nob_Latn",
    "oc": "oci_Latn", "pl": "pol_Latn", "pt": "por_Latn", "ro": "ron_Latn",
    "ru": "rus_Cyrl", "gd": "gla_Latn", "sr": "srp_Cyrl",
    "sk": "slk_Latn", "sl": "slv_Latn", "es": "spa_Latn", "sv": "swe_Latn",
    "tr": "tur_Latn", "uk": "ukr_Cyrl", "zh": "zho_Hans",
}

MAX_TOKENS = 150  # Über ~150 Tokens pro Chunk fällt die Qualität von NLLB-600M deutlich ab


# ── Textextraktion aus Dokumenten ────────────────────────────────────────────

def extract_text(input_path: str, soffice_path: str = "soffice") -> list[str]:
    """Liefert Textblöcke aus der Eingabedatei (formatabhängig)."""
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".txt":
        return _read_txt(input_path)
    if ext == ".pdf":
        return _extract_pdf_blocks(input_path)
    if ext == ".docx":
        try:
            return _read_docx(input_path)
        except ImportError:
            pass  # python-docx nicht installiert → Fallback auf LibreOffice
    return _convert_via_libreoffice(input_path, soffice_path)


def _read_txt(path: str) -> list[str]:
    # Encodings der Reihe nach versuchen: UTF-8 mit BOM, reines UTF-8, Windows-1252, Latin-1
    content = None
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
    if content is None:
        # Letzter Ausweg: nicht dekodierbare Bytes ersetzen statt abzubrechen
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
    # In Absätze splitten (Leerzeile als Trenner); zu kurze Fragmente verwerfen
    paras = [p.strip() for p in content.split("\n\n") if len(p.strip()) >= 15]
    return paras or ([content.strip()] if content.strip() else [])


def _extract_pdf_blocks(pdf_path: str) -> list[str]:
    """Liefert nicht-triviale Textblöcke aus einem PDF, einen String pro Block."""
    doc = fitz.Document(pdf_path)
    blocks = []
    for page in doc:
        for block in page.get_text("blocks"):
            # block: (x0, y0, x1, y1, text, block_no, block_type)
            if block[6] != 0:   # Bildblöcke überspringen
                continue
            text = block[4].strip()
            if len(text) < 15:  # Seitenzahlen, Kopfzeilen, einzelne Ziffern überspringen
                continue
            blocks.append(text)
    doc.close()
    return blocks


def _read_docx(path: str) -> list[str]:
    """Absätze direkt aus einer DOCX-Datei via python-docx lesen (kein Encoding-Verlust)."""
    if not _HAS_PYTHON_DOCX:
        raise ImportError("python-docx nicht installiert (pip install python-docx)")
    doc = _DocxDocument(path)
    paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) >= 15]
    return paras


def _convert_via_libreoffice(input_path: str, soffice_path: str = "soffice") -> list[str]:
    """Office-Dokument via LibreOffice headless nach TXT konvertieren und dann einlesen."""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            result = subprocess.run(
                [soffice_path, "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, input_path],
                capture_output=True, text=True, timeout=120,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                f"LibreOffice-Konvertierung hat das Zeitlimit überschritten (>120s) für "
                f"{os.path.basename(input_path)}"
            )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice-Konvertierung fehlgeschlagen (Exit {result.returncode}): "
                f"{result.stderr.strip()}"
            )
        txt_files = [f for f in os.listdir(tmpdir) if f.endswith(".txt")]
        if not txt_files:
            raise RuntimeError("LibreOffice hat keine .txt-Ausgabe erzeugt")
        return _read_txt(os.path.join(tmpdir, txt_files[0]))


# ── Übersetzung ───────────────────────────────────────────────────────────────

def translate_blocks(
    blocks: list[str],
    src_lang: str,
    tgt_lang: str,
    model_path: str,
    device: str,
    compute_type: str,
    beam_size: int,
    max_batch_size: int = 16,
) -> list[str]:
    """Übersetzt eine Liste von Textblöcken direkt via SentencePiece + CTranslate2.

    Optimierung: Alle Chunks aller Blöcke werden gesammelt und in EINEM
    translate_batch-Aufruf übersetzt (statt eines Modellaufrufs pro Chunk).
    CTranslate2 parallelisiert das intern (gesteuert über max_batch_size) — auf
    der GPU bringt das ein Vielfaches an Durchsatz, auf der CPU profitiert es vom
    Threading. Die Chunk→Block-Zuordnung wird gemerkt, damit die Übersetzungen
    danach korrekt wieder zu Blöcken zusammengesetzt werden.
    """
    import ctranslate2
    import sentencepiece as spm

    src_code = LANG_MAP[src_lang]
    tgt_code = LANG_MAP[tgt_lang]

    gpu_temp_suffix = f" (GPU-Temp={gpu_temperature()}°C)" if device.lower() == "cuda" else ""
    print(f"Lade Modell aus {model_path}…{gpu_temp_suffix}", flush=True)
    sp = spm.SentencePieceProcessor(os.path.join(model_path, "sentencepiece.bpe.model"))
    translator = ctranslate2.Translator(model_path, device=device, compute_type=compute_type)

    # 1) Alle Blöcke in Chunks zerlegen und tokenisierte Quell-Sequenzen sammeln.
    #    chunk_counts merkt sich, wie viele Chunks jeder Block beigesteuert hat.
    source_batch: list[list[str]] = []
    chunk_counts: list[int] = []
    for i, block in enumerate(blocks, 1):
        chunks = _chunk_block(block, sp, MAX_TOKENS)
        chunk_counts.append(len(chunks))
        for chunk in chunks:
            pieces = sp.encode(chunk, out_type=str)
            # NLLB-Eingabeformat: pieces + </s> + src_lang (entspricht NllbTokenizer)
            source_batch.append(pieces + ['</s>', src_code])
        print(f"  Block {i}/{len(blocks)} ({len(block)} Zeichen) → {len(chunks)} Chunk(s)",
              flush=True)

    if not source_batch:
        return ["" for _ in blocks]

    # 2) Alles in einem Aufruf übersetzen; CTranslate2 bildet intern Teilbatches.
    print(f"Übersetze {len(source_batch)} Chunk(s) in Batches "
          f"(max_batch_size={max_batch_size})…", flush=True)
    results = translator.translate_batch(
        source_batch,
        target_prefix=[[tgt_code]] * len(source_batch),
        beam_size=beam_size,
        repetition_penalty=1.2,
        no_repeat_ngram_size=4,
        max_batch_size=max_batch_size,
    )

    # 3) Hypothesen dekodieren (führendes Ziel-Sprachtoken und Sondertokens entfernen).
    decoded: list[str] = []
    for result in results:
        hypothesis = result.hypotheses[0]
        output_pieces = [t for t in hypothesis if t not in (tgt_code, '</s>', '<pad>')]
        decoded.append(sp.decode(output_pieces))

    # 4) Dekodierte Chunks anhand chunk_counts wieder den Blöcken zuordnen.
    translated: list[str] = []
    pos = 0
    for count in chunk_counts:
        translated.append(" ".join(decoded[pos:pos + count]))
        pos += count

    return translated


# Gängige deutsche/englische Abkürzungen, die mit '.' enden, aber keinen Satz beenden
_ABBREVS = {
    'dr', 'prof', 'hr', 'fr', 'st', 'nr', 'str', 'pl', 'abs', 'art',
    'bzw', 'ca', 'etc', 'evtl', 'ggf', 'inkl', 'max', 'min', 'mrd',
    'mio', 'rd', 'sog', 'usw', 'vgl', 'vs', 'z', 'b', 'd', 'h', 'o',
    'a', 'u', 's', 'p', 'i', 'e', 'fig', 'abb', 'tab', 'vol', 'ed',
    'pp', 'op', 'mr', 'ms', 'mrs', 'jan', 'feb', 'mär', 'apr', 'jun',
    'jul', 'aug', 'sep', 'okt', 'nov', 'dez',
}

_WHITESPACE_RE = re.compile(r'\s+')


def _ends_sentence(bare: str, next_word: str) -> bool:
    """Entscheidet, ob ein Wort einen Satz beendet.

    bare: Wort ohne nachgestellte Anführungszeichen/Klammern.
    next_word: nachfolgendes Wort ("" am Textende) — für die Punkt-Heuristik.
    """
    if bare.endswith(('!', '?', ';')):
        return True
    if not bare.endswith('.'):
        return False
    base = bare.rstrip('.').lower()
    # Abkürzungen sind kein Satzende: einzelnes Zeichen oder bekannte Abkürzung
    if len(base) <= 1 or base in _ABBREVS:
        return False
    # Punkt trennt nur, wenn das nächste Wort mit Großbuchstaben beginnt
    return bool(next_word) and next_word[0].isupper()


def _split_sentences(text: str) -> list[str]:
    """Zerlegt Text in Sätze; berücksichtigt deutsche Abkürzungen und Semikola."""
    # Whitespace normalisieren
    text = _WHITESPACE_RE.sub(' ', text).strip()
    if not text:
        return []

    result, current = [], []
    words = text.split(' ')
    for i, word in enumerate(words):
        current.append(word)
        bare = word.rstrip('"\')])')   # nachgestellte Anführungszeichen/Klammern entfernen
        next_word = words[i + 1] if i + 1 < len(words) else ""
        if _ends_sentence(bare, next_word):
            result.append(' '.join(current).strip())
            current = []
    if current:
        result.append(' '.join(current).strip())

    return [s for s in result if s] or [text]


def _chunk_block(text: str, sp, max_tokens: int) -> list[str]:
    """Fasst Sätze zu token-begrenzten Chunks zusammen; zu lange Sätze werden wortweise geteilt."""
    sentences = _split_sentences(text)
    if not sentences:
        return [text] if text.strip() else []

    chunks, current, current_len = [], [], 0
    for sentence in sentences:
        sentence_len = len(sp.encode(sentence))
        if sentence_len > max_tokens:
            # Satz allein schon zu lang → laufenden Chunk abschließen und wortweise splitten
            if current:
                chunks.append(' '.join(current))
                current, current_len = [], 0
            chunks.extend(_split_by_words(sentence, sp, max_tokens))
        elif current and current_len + sentence_len > max_tokens:
            # Limit würde überschritten → aktuellen Chunk abschließen, neuen beginnen
            chunks.append(' '.join(current))
            current, current_len = [sentence], sentence_len
        else:
            current.append(sentence)
            current_len += sentence_len
    if current:
        chunks.append(' '.join(current))
    return chunks or [text]


def _split_by_words(text: str, sp, max_tokens: int) -> list[str]:
    """Notfall-Aufteilung: Wörter gierig packen, bis der Chunk max_tokens überschreiten würde."""
    words = text.split()
    chunks, current, current_len = [], [], 0
    for word in words:
        word_len = len(sp.encode(word))
        if current and current_len + word_len > max_tokens:
            chunks.append(' '.join(current))
            current, current_len = [word], word_len
        else:
            current.append(word)
            current_len += word_len
    if current:
        chunks.append(' '.join(current))
    return chunks or [text]


# ── Einstiegspunkt ─────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input",          required=True)
    parser.add_argument("--output",         required=True)
    parser.add_argument("--source",         required=True)
    parser.add_argument("--target",         required=True)
    parser.add_argument("--model",          required=True)
    parser.add_argument("--libreoffice",    default="soffice", dest="libreoffice")
    parser.add_argument("--device",         default="cpu")
    parser.add_argument("--compute-type",   default="int8", dest="compute_type")
    parser.add_argument("--beam-size",      default=4, type=int, dest="beam_size")
    # Anzahl gleichzeitig übersetzter Chunks pro internem CTranslate2-Batch
    parser.add_argument("--max-batch-size", default=16, type=int, dest="max_batch_size")
    args = parser.parse_args()

    # Quell- und Zielsprache gegen die Sprachtabelle prüfen
    if args.source not in LANG_MAP:
        print(f"FEHLER: unbekannte Quellsprache: {args.source}", file=sys.stderr)
        return 1
    if args.target not in LANG_MAP:
        print(f"FEHLER: unbekannte Zielsprache: {args.target}", file=sys.stderr)
        return 1

    print(f"Extrahiere Text aus {args.input}…", flush=True)
    try:
        blocks = extract_text(args.input, args.libreoffice)
    except Exception as e:
        print(f"FEHLER beim Extrahieren des Texts: {e}", file=sys.stderr)
        return 1
    print(f"  {len(blocks)} Blöcke extrahiert", flush=True)
    for idx, b in enumerate(blocks[:5]):
        print(f"  block[{idx}] ({len(b)} Zeichen): {b[:100]!r}", flush=True)

    if not blocks:
        print("Keine Textblöcke gefunden — schreibe leere Ausgabe.", flush=True)
        open(args.output, "w").close()
        return 0

    try:
        translated = translate_blocks(
            blocks,
            src_lang=args.source,
            tgt_lang=args.target,
            model_path=args.model,
            device=args.device,
            compute_type=args.compute_type,
            beam_size=args.beam_size,
            max_batch_size=args.max_batch_size,
        )
    except Exception as e:
        print(f"FEHLER während der Übersetzung: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1

    # Atomar schreiben (tmp-Datei + os.replace): nach einem potenziell langen
    # Batch-Übersetzungslauf soll ein Absturz mitten im Schreiben nicht eine
    # unvollständige, aber vorhandene Output-Datei hinterlassen.
    tmp_output = args.output + ".tmp"
    with open(tmp_output, "w", encoding="utf-8") as f:
        f.write("\n\n".join(translated))
    os.replace(tmp_output, args.output)

    gpu_temp_suffix = f" (GPU-Temp={gpu_temperature()}°C)" if args.device.lower() == "cuda" else ""
    print(f"Fertig → {args.output}{gpu_temp_suffix}", flush=True)
    return 0


if __name__ == "__main__":
    sys.exit(main())
